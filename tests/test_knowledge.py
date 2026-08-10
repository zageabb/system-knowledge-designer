from io import BytesIO
from email.message import EmailMessage
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document as DocxDocument
from openpyxl import Workbook

import services.knowledge as knowledge_service
from app import create_app
from database import db
from models import ChunkKnowledgeLink, DocumentChunk, DocumentVersion, EvidenceRecord, KnowledgeDocument, KnowledgeLink, SystemProject
from services.er_language import parse_er_source
from services.revisions import create_revision
from tests.test_sandbox_sql import prepare as prepare_model


def make_app(tmp_path):
    return create_app({"TESTING": True, "SECRET_KEY": "test", "SQLALCHEMY_DATABASE_URI": f"sqlite:///{tmp_path/'knowledge.db'}", "DATA_DIR": tmp_path/"data", "WTF_CSRF_ENABLED": False, "ADMIN_PASSWORD": "test-password"})


def prepare(app):
    with app.app_context():
        project = SystemProject(name="Procurement", slug="procurement", dialect="sqlite")
        db.session.add(project); db.session.commit(); return project.id


def test_upload_extract_index_search_and_citation(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={
        "document": (BytesIO(b"# Invoice Approval\n\nInvoices above 5000 require finance approval."), "approval.md"),
        "title": "Approval policy", "classification": "internal", "provenance": "authored",
    }, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=finance+approval")
    assert search.status_code == 200 and b"Approval policy" in search.data and b"Heading: Invoice Approval" in search.data
    with app.app_context():
        document = KnowledgeDocument.query.one(); chunk = DocumentChunk.query.filter(DocumentChunk.text.contains("finance approval")).one()
        assert document.managed_path.startswith(str((tmp_path / "data").resolve()))
        chunk_id = chunk.id
    citation = client.get(f"/projects/{project_id}/knowledge/chunks/{chunk_id}")
    assert citation.status_code == 200 and b"Invoices above 5000" in citation.data


def test_document_title_links_to_project_scoped_viewer(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"First section\n\nSecond section"), "viewer.txt"), "title": "Viewer document"}, content_type="multipart/form-data", follow_redirects=True)
    with app.app_context(): document_id = KnowledgeDocument.query.one().id
    viewer_path = f"/projects/{project_id}/knowledge/documents/{document_id}"
    assert response.status_code == 200 and f'href="{viewer_path}"'.encode() in response.data
    viewer = client.get(viewer_path)
    assert viewer.status_code == 200 and b"MANAGED DOCUMENT VIEWER" in viewer.data and b"First section" in viewer.data
    assert b"Open citation and model links" in viewer.data


def test_document_viewer_rejects_cross_project_document(tmp_path):
    app = make_app(tmp_path); first_id = prepare(app)
    with app.app_context():
        second = SystemProject(name="Finance", slug="finance", dialect="sqlite"); db.session.add(second); db.session.commit(); second_id = second.id
    client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{first_id}/knowledge", data={"document": (BytesIO(b"Private viewer content"), "private.txt")}, content_type="multipart/form-data")
    with app.app_context(): document_id = KnowledgeDocument.query.one().id
    rejected = client.get(f"/projects/{second_id}/knowledge/documents/{document_id}")
    assert rejected.status_code == 400 and b"does not belong to project" in rejected.data


def test_unsupported_upload_is_rejected_without_record(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"binary"), "unsafe.exe")}, content_type="multipart/form-data", follow_redirects=True)
    assert b"Supported document formats" in response.data
    with app.app_context(): assert KnowledgeDocument.query.count() == 0


def test_citation_is_project_scoped(tmp_path):
    app = make_app(tmp_path); first_id = prepare(app)
    with app.app_context():
        second = SystemProject(name="Finance", slug="finance", dialect="sqlite"); db.session.add(second); db.session.commit(); second_id = second.id
    client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{first_id}/knowledge", data={"document": (BytesIO(b"Private project evidence"), "note.txt")}, content_type="multipart/form-data")
    with app.app_context(): chunk_id = DocumentChunk.query.one().id
    assert client.get(f"/projects/{second_id}/knowledge/chunks/{chunk_id}").status_code == 400


def test_federated_search_returns_schema_relationship_and_sample_evidence(tmp_path):
    app = make_app(tmp_path); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    project_id, revision_id, dataset_id = prepare_model(app)
    response = client.get(f"/projects/{project_id}/knowledge?q=supplier_id")
    assert response.status_code == 200
    assert b"Authoritative schema" in response.data
    assert b"PURCHASE_ORDER.supplier_id" in response.data
    assert b"Representative sample values" in response.data
    assert b"supplier_id=1" in response.data


def test_federated_search_source_filter_limits_families(tmp_path):
    app = make_app(tmp_path); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    project_id, revision_id, dataset_id = prepare_model(app)
    response = client.get(f"/projects/{project_id}/knowledge?q=supplier_id&source=schema")
    assert b"Authoritative schema" in response.data
    assert b"Representative sample values" not in response.data


def test_csv_upload_uses_row_range_citations(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={
        "document": (BytesIO(b"supplier_id,supplier_name\n101,Synthetic Components\n102,Example Logistics\n"), "suppliers.csv"),
        "title": "Supplier extract",
    }, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=Logistics&source=documents")
    assert b"Rows 2" in search.data and b"supplier_name=Example [Logistics]" in search.data


def test_delete_document_removes_chunks_index_and_managed_file(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Unique deletion evidence"), "delete-me.txt")}, content_type="multipart/form-data")
    with app.app_context():
        document = KnowledgeDocument.query.one(); document_id = document.id
        managed_path = document.managed_path
    response = client.post(f"/projects/{project_id}/knowledge/documents/{document_id}/delete", follow_redirects=True)
    assert response.status_code == 200 and b"document &#39;delete-me.txt&#39; deleted" in response.data
    with app.app_context():
        assert db.session.get(KnowledgeDocument, document_id) is None
        assert DocumentChunk.query.count() == 0
    assert not __import__("pathlib").Path(managed_path).exists()
    search = client.get(f"/projects/{project_id}/knowledge?q=deletion&source=documents")
    assert b"No matching local evidence" in search.data


def test_docx_upload_preserves_heading_locator(tmp_path):
    source = BytesIO(); document = DocxDocument(); document.add_heading("Invoice Controls", level=1); document.add_paragraph("Three-way matching is mandatory."); document.save(source); source.seek(0)
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (source, "controls.docx")}, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=mandatory&source=documents")
    assert b"Heading: Invoice Controls" in search.data and b"Three-way matching" in search.data


def test_xlsx_upload_preserves_sheet_and_row_locator(tmp_path):
    source = BytesIO(); workbook = Workbook(); sheet = workbook.active; sheet.title = "Suppliers"; sheet.append(["supplier_id", "supplier_name"]); sheet.append([101, "Synthetic Metals"]); workbook.save(source); source.seek(0)
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (source, "suppliers.xlsx")}, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=Metals&source=documents")
    assert b"Sheet: Suppliers, rows 2" in search.data and b"supplier_id=101" in search.data


def test_pdf_extractor_uses_page_locators(monkeypatch):
    class Page:
        def __init__(self, value): self.value = value
        def extract_text(self): return self.value
    class Reader:
        def __init__(self, stream): self.pages = [Page("Approval overview"), Page("Finance authorisation")]
    monkeypatch.setattr(knowledge_service, "PdfReader", Reader)
    chunks = knowledge_service.extract_structured_document(b"%PDF-synthetic", ".pdf")
    assert chunks == [("Page 1", "Approval overview"), ("Page 2", "Finance authorisation")]


def test_invalid_office_container_is_rejected():
    with __import__("pytest").raises(knowledge_service.KnowledgeIngestionError, match="valid ZIP package"):
        knowledge_service.extract_structured_document(b"not-a-zip", ".docx")


def test_eml_upload_preserves_message_locator(tmp_path):
    message = EmailMessage(); message["Subject"] = "Invoice escalation"; message["Date"] = "Fri, 31 Jul 2026 10:00:00 +0100"; message.set_content("Escalate overdue invoices to finance operations.")
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(message.as_bytes()), "approval.eml")}, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=overdue&source=documents")
    assert b"Message:" in search.data and b"Invoice escalation" in search.data and b"finance operations" in search.data


def test_zip_upload_indexes_supported_members_with_file_locator(tmp_path):
    source = BytesIO()
    with ZipFile(source, "w", ZIP_DEFLATED) as archive:
        archive.writestr("nested/rules.md", "# Payment Rules\n\nDual approval is required.")
        archive.writestr("ignored.bin", b"\x00\x01")
    source.seek(0)
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (source, "knowledge.zip")}, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"searchable chunk" in response.data
    search = client.get(f"/projects/{project_id}/knowledge?q=approval&source=documents")
    assert b"File: rules.md" in search.data and b"Dual [approval]" in search.data


def test_zip_without_supported_members_is_rejected(tmp_path):
    source = BytesIO()
    with ZipFile(source, "w", ZIP_DEFLATED) as archive: archive.writestr("image.bin", b"binary")
    source.seek(0)
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.post(f"/projects/{project_id}/knowledge", data={"document": (source, "empty.zip")}, content_type="multipart/form-data", follow_redirects=True)
    assert b"no supported searchable" in response.data
    with app.app_context(): assert KnowledgeDocument.query.count() == 0


def test_user_can_create_and_remove_validated_document_model_link(tmp_path):
    app = make_app(tmp_path); project_id, revision_id, dataset_id = prepare_model(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Supplier operating guidance"), "supplier.txt")}, content_type="multipart/form-data")
    with app.app_context(): document_id = KnowledgeDocument.query.one().id
    created = client.post(f"/projects/{project_id}/knowledge/documents/{document_id}/links", data={"target": "column|SUPPLIER.supplier_id"}, follow_redirects=True)
    assert created.status_code == 200 and b"Linked &#39;supplier.txt&#39; to SUPPLIER.supplier_id" in created.data
    with app.app_context():
        link = KnowledgeLink.query.one(); link_id = link.id
        assert link.revision_id == revision_id and link.created_by == "admin"
        chunk_id = DocumentChunk.query.one().id
    citation = client.get(f"/projects/{project_id}/knowledge/chunks/{chunk_id}")
    assert b"Linked model objects" in citation.data and b"SUPPLIER.supplier_id" in citation.data
    removed = client.post(f"/projects/{project_id}/knowledge/links/{link_id}/delete", follow_redirects=True)
    assert removed.status_code == 200 and b"Removed document link" in removed.data
    with app.app_context(): assert KnowledgeLink.query.count() == 0


def test_document_model_link_rejects_unknown_target(tmp_path):
    app = make_app(tmp_path); project_id, revision_id, dataset_id = prepare_model(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Evidence"), "note.txt")}, content_type="multipart/form-data")
    with app.app_context(): document_id = KnowledgeDocument.query.one().id
    rejected = client.post(f"/projects/{project_id}/knowledge/documents/{document_id}/links", data={"target": "column|SUPPLIER.missing"}, follow_redirects=True)
    assert b"not valid for the active revision" in rejected.data
    with app.app_context(): assert KnowledgeLink.query.count() == 0


def test_search_returns_bounded_multi_hop_relationship_path(tmp_path):
    source = '''erModel Routes {
      subjectArea Core {
        table SUPPLIER {
          integer supplier_id PK
        }
        table PURCHASE_ORDER {
          integer purchase_order_id PK
          integer supplier_id FK
        }
        table INVOICE {
          integer invoice_id PK
          integer purchase_order_id FK
        }
        relationship PURCHASE_ORDER.supplier_id -> SUPPLIER.supplier_id {
          cardinality many-to-one
        }
        relationship INVOICE.purchase_order_id -> PURCHASE_ORDER.purchase_order_id {
          cardinality many-to-one
        }
      }
    }'''
    app = make_app(tmp_path)
    with app.app_context():
        project = SystemProject(name="Routes", slug="routes", dialect="sqlite"); db.session.add(project); db.session.flush()
        revision = create_revision(project, source, parse_er_source(source), "approved"); revision.status = "approved"; project.active_revision_id = revision.id; db.session.commit(); project_id = project.id
    client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    response = client.get(f"/projects/{project_id}/knowledge?q=supplier+invoice&source=schema")
    assert response.status_code == 200 and b"Relationship path: SUPPLIER to INVOICE" in response.data
    assert b"2 hops" in response.data and b"PURCHASE_ORDER.supplier_id" in response.data and b"INVOICE.purchase_order_id" in response.data


def test_upload_new_document_version_preserves_prior_version_and_citations(tmp_path):
    app = make_app(tmp_path); project_id = prepare(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Original approval threshold is 5000."), "policy-v1.txt"), "title": "Approval policy"}, content_type="multipart/form-data")
    with app.app_context(): first = KnowledgeDocument.query.one(); first_id = first.id; family_id = first.version.family_id
    response = client.post(f"/projects/{project_id}/knowledge/documents/{first_id}/versions", data={"document": (BytesIO(b"Revised approval threshold is 7500."), "policy-v2.txt")}, content_type="multipart/form-data", follow_redirects=True)
    assert response.status_code == 200 and b"Uploaded version 2" in response.data
    with app.app_context():
        versions = DocumentVersion.query.filter_by(family_id=family_id).order_by(DocumentVersion.version_number).all()
        assert [version.version_number for version in versions] == [1, 2]
        assert versions[1].predecessor_document_id == first_id
        assert KnowledgeDocument.query.count() == 2 and DocumentChunk.query.count() == 2
    assert b"[Original] approval" in client.get(f"/projects/{project_id}/knowledge?q=Original&source=documents").data
    assert b"[Revised] approval" in client.get(f"/projects/{project_id}/knowledge?q=Revised&source=documents").data


def test_knowledge_coverage_and_stale_link_diagnostics(tmp_path):
    app = make_app(tmp_path); project_id, revision_id, dataset_id = prepare_model(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Supplier guidance"), "supplier.txt")}, content_type="multipart/form-data")
    with app.app_context(): document_id = KnowledgeDocument.query.one().id
    unlinked = client.get(f"/projects/{project_id}/knowledge")
    assert b"0 / 2" in unlinked.data and b"Documents without model links" in unlinked.data
    client.post(f"/projects/{project_id}/knowledge/documents/{document_id}/links", data={"target": "table|SUPPLIER"})
    linked = client.get(f"/projects/{project_id}/knowledge")
    assert b"1 / 2" in linked.data
    replacement = '''erModel Replacement {
      subjectArea Core {
        table INVOICE {
          integer invoice_id PK
        }
      }
    }'''
    with app.app_context():
        project = db.session.get(SystemProject, project_id); revision = create_revision(project, replacement, parse_er_source(replacement), "replacement"); revision.status = "approved"; project.active_revision_id = revision.id; db.session.commit()
    stale = client.get(f"/projects/{project_id}/knowledge")
    assert b"Stale model links" in stale.data and b"SUPPLIER" in stale.data and b"do not exist in the active approved model" in stale.data


def test_user_can_link_specific_cited_passage_and_coverage_counts_it(tmp_path):
    app = make_app(tmp_path); project_id, revision_id, dataset_id = prepare_model(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    client.post(f"/projects/{project_id}/knowledge", data={"document": (BytesIO(b"Supplier identifier definition"), "definition.txt")}, content_type="multipart/form-data")
    with app.app_context(): chunk_id = DocumentChunk.query.one().id
    created = client.post(f"/projects/{project_id}/knowledge/chunks/{chunk_id}/links", data={"target": "column|SUPPLIER.supplier_id"}, follow_redirects=True)
    assert created.status_code == 200 and b"Linked cited passage" in created.data and b"This cited passage" in created.data
    with app.app_context():
        link = ChunkKnowledgeLink.query.one(); link_id = link.id
        assert link.revision_id == revision_id and link.created_by == "admin"
    coverage = client.get(f"/projects/{project_id}/knowledge")
    assert b"1 / 5" in coverage.data and b"Documents without model links" in coverage.data
    removed = client.post(f"/projects/{project_id}/knowledge/chunk-links/{link_id}/delete", follow_redirects=True)
    assert b"Removed passage link" in removed.data
    with app.app_context(): assert ChunkKnowledgeLink.query.count() == 0


def test_user_can_save_open_and_delete_evidence_snapshot(tmp_path):
    app = make_app(tmp_path); project_id, revision_id, dataset_id = prepare_model(app); client = app.test_client(); client.post("/login", data={"username": "admin", "password": "test-password"})
    saved = client.post(f"/projects/{project_id}/knowledge/evidence", data={"q": "supplier_id", "source": "all"}, follow_redirects=True)
    assert saved.status_code == 200 and b"PERSISTED LOCAL EVIDENCE" in saved.data and b"PURCHASE_ORDER.supplier_id" in saved.data
    with app.app_context():
        record = EvidenceRecord.query.one(); record_id = record.id; payload = __import__("json").loads(record.evidence_json)
        assert record.model_revision_id == revision_id and record.created_by == "admin"
        assert payload["schema"] and payload["relationships"] and payload["samples"]
    listing = client.get(f"/projects/{project_id}/knowledge")
    assert b"Saved evidence snapshots" in listing.data and f"#{record_id}".encode() in listing.data
    deleted = client.post(f"/projects/{project_id}/knowledge/evidence/{record_id}/delete", follow_redirects=True)
    assert b"Evidence snapshot" in deleted.data and b"deleted" in deleted.data
    with app.app_context(): assert EvidenceRecord.query.count() == 0
