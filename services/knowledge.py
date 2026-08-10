from __future__ import annotations

import hashlib
import csv
import io
import re
import zipfile
from email import policy
from email.parser import BytesParser
from pathlib import Path
from uuid import uuid4

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from sqlalchemy import text

from database import db
from models import DocumentChunk, KnowledgeDocument


ALLOWED_EXTENSIONS = {".txt": "text/plain", ".md": "text/markdown", ".markdown": "text/markdown", ".csv": "text/csv", ".pdf": "application/pdf", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".eml": "message/rfc822", ".zip": "application/zip"}
MAX_EXTRACTED_CHARS = 2_000_000
MAX_PDF_PAGES = 500
MAX_WORKBOOK_CELLS = 100_000
MAX_ARCHIVE_ENTRIES = 5_000
MAX_ARCHIVE_EXPANDED_BYTES = 50 * 1024 * 1024


class KnowledgeIngestionError(ValueError):
    pass


def ensure_fts_index() -> None:
    db.session.execute(text("CREATE VIRTUAL TABLE IF NOT EXISTS document_chunk_fts USING fts5(chunk_id UNINDEXED, project_id UNINDEXED, title, locator, content)"))
    db.session.commit()


def ingest_document(*, project_id: int, upload, title: str, provenance: str, classification: str, data_dir: Path) -> KnowledgeDocument:
    original = Path(upload.filename or "").name
    extension = Path(original).suffix.casefold()
    if extension not in ALLOWED_EXTENSIONS:
        raise KnowledgeIngestionError("Supported document formats are TXT, Markdown, CSV, PDF, DOCX, XLSX, EML and controlled ZIP.")
    payload = upload.read()
    if not payload:
        raise KnowledgeIngestionError("The uploaded document is empty.")
    if extension in {".pdf", ".docx", ".xlsx", ".eml", ".zip"}:
        chunks = extract_structured_document(payload, extension)
    else:
        try:
            content = payload.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise KnowledgeIngestionError("TXT, Markdown and CSV uploads must use UTF-8 encoding.") from exc
        chunks = chunk_csv(content) if extension == ".csv" else chunk_text(content)
    if not chunks:
        raise KnowledgeIngestionError("The uploaded document contains no searchable text.")
    root = Path(data_dir).resolve() / str(project_id) / "documents"
    root.mkdir(parents=True, exist_ok=True)
    managed_path = root / f"{uuid4().hex}{extension}"
    managed_path.write_bytes(payload)
    document = KnowledgeDocument(project_id=project_id, title=(title.strip() or original), original_filename=original, media_type=ALLOWED_EXTENSIONS[extension], managed_path=str(managed_path), content_hash=hashlib.sha256(payload).hexdigest(), provenance=provenance, classification=classification)
    db.session.add(document); db.session.flush()
    for position, (locator, chunk) in enumerate(chunks, 1):
        item = DocumentChunk(document_id=document.id, position=position, locator=locator, text=chunk)
        db.session.add(item); db.session.flush()
        db.session.execute(text("INSERT INTO document_chunk_fts(chunk_id, project_id, title, locator, content) VALUES (:chunk_id, :project_id, :title, :locator, :content)"), {"chunk_id": item.id, "project_id": project_id, "title": document.title, "locator": locator, "content": chunk})
    return document


def ingest_external_citation(*, project_id: int, job_id: int, citation_index: int, title: str, url: str, excerpt: str, data_dir: Path) -> KnowledgeDocument:
    from urllib.parse import parse_qs, urlparse

    parsed = urlparse(url)
    page_ids = parse_qs(parsed.query).get("curid", [])
    if parsed.scheme != "https" or parsed.netloc != "en.wikipedia.org" or parsed.path != "/" or len(page_ids) != 1 or not page_ids[0].isdigit():
        raise KnowledgeIngestionError("The external citation URL is not an approved Wikipedia page reference.")
    clean_title = title.strip()[:240]
    clean_excerpt = excerpt.strip()[:1000]
    if not clean_title or not clean_excerpt:
        raise KnowledgeIngestionError("The external citation has no promotable title or excerpt.")
    content = f"# {clean_title}\n\n{clean_excerpt}\n\nSource: {url}\n"
    payload = content.encode("utf-8")
    root = Path(data_dir).resolve() / str(project_id) / "documents"
    root.mkdir(parents=True, exist_ok=True)
    original = f"wikipedia-job-{job_id}-citation-{citation_index + 1}.md"
    managed_path = root / f"external-{job_id}-{citation_index}.md"
    if managed_path.exists():
        raise KnowledgeIngestionError("This external citation already has a managed source file.")
    managed_path.write_bytes(payload)
    document = KnowledgeDocument(project_id=project_id, title=clean_title, original_filename=original, media_type="text/markdown", managed_path=str(managed_path), content_hash=hashlib.sha256(payload).hexdigest(), provenance="external-wikipedia", classification="public")
    db.session.add(document); db.session.flush()
    item = DocumentChunk(document_id=document.id, position=1, locator=f"Wikipedia citation · {url}", text=content)
    db.session.add(item); db.session.flush()
    db.session.execute(text("INSERT INTO document_chunk_fts(chunk_id, project_id, title, locator, content) VALUES (:chunk_id, :project_id, :title, :locator, :content)"), {"chunk_id": item.id, "project_id": project_id, "title": document.title, "locator": item.locator, "content": item.text})
    return document


def chunk_text(content: str, max_chars: int = 1400) -> list[tuple[str, str]]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", content.replace("\r\n", "\n")) if block.strip()]
    result = []
    heading = ""
    for block in blocks:
        if block.startswith("#"):
            heading = block.lstrip("#").strip()
        for start in range(0, len(block), max_chars):
            text_value = block[start:start + max_chars].strip()
            if text_value:
                locator = f"Heading: {heading}" if heading else f"Chunk {len(result) + 1}"
                result.append((locator, text_value))
    return result


def chunk_csv(content: str, rows_per_chunk: int = 25) -> list[tuple[str, str]]:
    try:
        rows = list(csv.reader(io.StringIO(content)))
    except csv.Error as exc:
        raise KnowledgeIngestionError(f"CSV extraction failed: {exc}") from exc
    if not rows or not any(cell.strip() for row in rows for cell in row):
        return []
    header = rows[0]
    data_rows = rows[1:]
    if not data_rows:
        return [("Header row", ", ".join(header))]
    chunks = []
    for start in range(0, len(data_rows), rows_per_chunk):
        batch = data_rows[start:start + rows_per_chunk]
        lines = [" | ".join(f"{header[index] if index < len(header) else f'column_{index + 1}'}={value}" for index, value in enumerate(row)) for row in batch]
        first_row = start + 2; last_row = start + len(batch) + 1
        chunks.append((f"Rows {first_row}–{last_row}", "\n".join(lines)))
    return chunks


def extract_structured_document(payload: bytes, extension: str) -> list[tuple[str, str]]:
    try:
        if extension == ".pdf":
            if not payload.startswith(b"%PDF-"): raise KnowledgeIngestionError("The file does not have a valid PDF signature.")
            return extract_pdf(payload)
        if extension in {".docx", ".xlsx"}: validate_office_container(payload)
        if extension == ".docx": return extract_docx(payload)
        if extension == ".xlsx": return extract_xlsx(payload)
        if extension == ".eml": return extract_eml(payload)
        if extension == ".zip": return extract_zip(payload)
    except KnowledgeIngestionError:
        raise
    except Exception as exc:
        raise KnowledgeIngestionError(f"{extension[1:].upper()} extraction failed: {exc}") from exc
    raise KnowledgeIngestionError(f"No extractor is configured for {extension}.")


def validate_office_container(payload: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            members = archive.infolist()
            expanded = sum(member.file_size for member in members)
            compressed = sum(member.compress_size for member in members)
    except zipfile.BadZipFile as exc:
        raise KnowledgeIngestionError("The Office document is not a valid ZIP package.") from exc
    if len(members) > MAX_ARCHIVE_ENTRIES: raise KnowledgeIngestionError("The Office document contains too many package entries.")
    if expanded > MAX_ARCHIVE_EXPANDED_BYTES: raise KnowledgeIngestionError("The expanded Office document exceeds the safety limit.")
    if compressed and expanded / compressed > 100: raise KnowledgeIngestionError("The Office document compression ratio exceeds the safety limit.")


def validate_archive(archive: zipfile.ZipFile) -> list[zipfile.ZipInfo]:
    members = [member for member in archive.infolist() if not member.is_dir()]
    expanded = sum(member.file_size for member in members); compressed = sum(member.compress_size for member in members)
    if len(members) > MAX_ARCHIVE_ENTRIES: raise KnowledgeIngestionError("The archive contains too many entries.")
    if expanded > MAX_ARCHIVE_EXPANDED_BYTES: raise KnowledgeIngestionError("The expanded archive exceeds the safety limit.")
    if compressed and expanded / compressed > 100: raise KnowledgeIngestionError("The archive compression ratio exceeds the safety limit.")
    return members


def extract_eml(payload: bytes) -> list[tuple[str, str]]:
    message = BytesParser(policy=policy.default).parsebytes(payload)
    subject = str(message.get("Subject", "No subject")); sent = str(message.get("Date", "Unknown date"))
    bodies = []
    parts = message.walk() if message.is_multipart() else [message]
    for part in parts:
        if part.get_content_type() == "text/plain" and part.get_content_disposition() != "attachment":
            try: bodies.append(part.get_content())
            except (LookupError, UnicodeDecodeError): continue
    content = "\n\n".join(bodies).strip()
    if len(content) > MAX_EXTRACTED_CHARS: raise KnowledgeIngestionError("Extracted message text exceeds the safety limit.")
    return [(f"Message: {sent} · {subject}", value) for _, value in chunk_text(content)]


def extract_zip(payload: bytes) -> list[tuple[str, str]]:
    try:
        archive = zipfile.ZipFile(io.BytesIO(payload)); members = validate_archive(archive)
    except zipfile.BadZipFile as exc:
        raise KnowledgeIngestionError("The ZIP upload is not a valid archive.") from exc
    chunks = []
    with archive:
        for member in members:
            member_name = Path(member.filename).name
            extension = Path(member_name).suffix.casefold()
            if extension not in {".txt", ".md", ".markdown", ".csv"}: continue
            try: content = archive.read(member).decode("utf-8-sig")
            except UnicodeDecodeError as exc: raise KnowledgeIngestionError(f"ZIP member {member_name} is not valid UTF-8.") from exc
            extracted = chunk_csv(content) if extension == ".csv" else chunk_text(content)
            chunks.extend((f"File: {member_name} · {locator}", value) for locator, value in extracted)
    if not chunks: raise KnowledgeIngestionError("The ZIP contains no supported searchable TXT, Markdown or CSV files.")
    return chunks


def extract_pdf(payload: bytes) -> list[tuple[str, str]]:
    reader = PdfReader(io.BytesIO(payload))
    if len(reader.pages) > MAX_PDF_PAGES: raise KnowledgeIngestionError(f"PDF exceeds the {MAX_PDF_PAGES}-page extraction limit.")
    chunks = []; total = 0
    for page_number, page in enumerate(reader.pages, 1):
        page_text = (page.extract_text() or "").strip(); total += len(page_text)
        if total > MAX_EXTRACTED_CHARS: raise KnowledgeIngestionError("Extracted document text exceeds the safety limit.")
        for part, start in enumerate(range(0, len(page_text), 1400), 1):
            value = page_text[start:start + 1400].strip()
            if value: chunks.append((f"Page {page_number}" + (f", part {part}" if len(page_text) > 1400 else ""), value))
    return chunks


def extract_docx(payload: bytes) -> list[tuple[str, str]]:
    document = Document(io.BytesIO(payload)); chunks = []; heading = ""; total = 0
    for paragraph_number, paragraph in enumerate(document.paragraphs, 1):
        value = paragraph.text.strip()
        if not value: continue
        total += len(value)
        if total > MAX_EXTRACTED_CHARS: raise KnowledgeIngestionError("Extracted document text exceeds the safety limit.")
        if paragraph.style and paragraph.style.name.casefold().startswith("heading"): heading = value
        locator = f"Heading: {heading}" if heading else f"Paragraph {paragraph_number}"
        for start in range(0, len(value), 1400): chunks.append((locator, value[start:start + 1400]))
    return chunks


def extract_xlsx(payload: bytes, rows_per_chunk: int = 25) -> list[tuple[str, str]]:
    workbook = load_workbook(io.BytesIO(payload), read_only=True, data_only=True); chunks = []; cell_count = 0
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            cell_count += len(row)
            if cell_count > MAX_WORKBOOK_CELLS: raise KnowledgeIngestionError(f"Workbook exceeds the {MAX_WORKBOOK_CELLS:,}-cell extraction limit.")
            rows.append(["" if value is None else str(value) for value in row])
        if not rows: continue
        header = rows[0]
        for start in range(0, len(rows[1:]), rows_per_chunk):
            batch = rows[1:][start:start + rows_per_chunk]
            lines = [" | ".join(f"{header[index] or f'column_{index + 1}'}={value}" for index, value in enumerate(row) if value) for row in batch]
            lines = [line for line in lines if line]
            if lines:
                first_row = start + 2; last_row = start + len(batch) + 1
                chunks.append((f"Sheet: {worksheet.title}, rows {first_row}–{last_row}", "\n".join(lines)))
    workbook.close(); return chunks


def remove_from_fts(chunk_ids: list[int]) -> None:
    for chunk_id in chunk_ids:
        db.session.execute(text("DELETE FROM document_chunk_fts WHERE chunk_id = :chunk_id"), {"chunk_id": chunk_id})


def search_documents(*, project_id: int, query: str, limit: int = 20) -> list[dict]:
    terms = re.findall(r"[\w-]+", query, flags=re.UNICODE)
    if not terms:
        return []
    match = " AND ".join(f'"{term}"' for term in terms[:10])
    rows = db.session.execute(text("""
        SELECT chunk_id, title, locator, snippet(document_chunk_fts, 4, '[', ']', ' … ', 24) AS excerpt
        FROM document_chunk_fts WHERE project_id = :project_id AND document_chunk_fts MATCH :match
        ORDER BY bm25(document_chunk_fts) LIMIT :limit
    """), {"project_id": project_id, "match": match, "limit": min(max(limit, 1), 50)}).mappings()
    return [dict(row) for row in rows]
